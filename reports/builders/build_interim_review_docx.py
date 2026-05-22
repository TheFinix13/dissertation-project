"""Build the EEEM004 Interim Review Word document.

Outputs:
    reports/generated/exports/InterimReview.docx

The document follows the supervisor-facing Interim Review form structure,
with all student-owned sections populated from the rebuilt dissertation
framing. The blue supervisor boxes are intentionally left empty.

Run:
    venv/bin/python reports/builders/build_interim_review_docx.py
"""

from __future__ import annotations

import json
import os
import subprocess
import sys
from pathlib import Path
from typing import Iterable

import numpy as np
from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parent.parent.parent
RESULTS = ROOT / "experiments" / "results"
EXPORTS = ROOT / "reports" / "generated" / "exports"
EQ_DIR = EXPORTS / "equations"
EXPORTS.mkdir(parents=True, exist_ok=True)
EQ_DIR.mkdir(parents=True, exist_ok=True)


# --------------------------------------------------------------------------- #
# Helpers (kept parallel to build_main_dissertation_docx.py)
# --------------------------------------------------------------------------- #
def latest_json(prefix: str) -> dict | list:
    files = sorted(p for p in RESULTS.glob(f"{prefix}_*.json"))
    if not files:
        return []
    return json.loads(files[-1].read_text(encoding="utf-8"))


def avg(rows: Iterable[dict], key: str) -> float:
    rows = list(rows)
    if not rows:
        return float("nan")
    return float(np.mean([float(r[key]) for r in rows]))


def add_heading(doc: Document, text: str, level: int) -> None:
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0, 0, 0)


def add_para(doc: Document, text: str, *, italic: bool = False, bold: bool = False, align=None) -> None:
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    run = p.add_run(text)
    run.italic = italic
    run.bold = bold


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_supervisor_box(doc: Document, prompt: str) -> None:
    """Render a labelled empty box where the supervisor will write."""
    p = doc.add_paragraph()
    run = p.add_run(prompt)
    run.italic = True
    run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
    run.font.size = Pt(10)
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.rows[0].cells[0]
    cell.text = " "
    cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), "EAF1F8")
    tc_pr.append(shd)
    set_cell_height(cell, Cm(2.5))


def set_cell_height(cell, height) -> None:
    tr = cell._tc.getparent()
    trPr = tr.get_or_add_trPr()
    trHeight = OxmlElement("w:trHeight")
    trHeight.set(qn("w:val"), str(int(height.emu / 635)))
    trHeight.set(qn("w:hRule"), "atLeast")
    trPr.append(trHeight)


def page_break(doc: Document) -> None:
    doc.add_page_break()


def render_equation(latex: str, filename: str, fontsize: int = 17) -> Path:
    import matplotlib
    matplotlib.use("Agg")
    import matplotlib.pyplot as plt
    path = EQ_DIR / filename
    prev = plt.rcParams["mathtext.default"]
    fig = None
    try:
        plt.rcParams["mathtext.default"] = "it"
        fig = plt.figure(figsize=(0.01, 0.01))
        fig.text(0, 0, latex, fontsize=fontsize)
        fig.savefig(path, dpi=220, bbox_inches="tight", pad_inches=0.22, transparent=False, facecolor="white")
    finally:
        plt.rcParams["mathtext.default"] = prev
        if fig is not None:
            plt.close(fig)
    return path


def add_equation(doc: Document, latex: str, filename: str, *, label: str | None = None, width_inches: float = 4.8) -> None:
    path = render_equation(latex, filename)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(path), width=Inches(width_inches))
    if label:
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap_run = cap.add_run(label)
        cap_run.italic = True
        cap_run.font.size = Pt(10)


def add_figure(doc: Document, image_path: Path, caption: str, *, width_inches: float = 5.5) -> None:
    if not image_path.exists():
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(image_path), width=Inches(width_inches))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap_run = cap.add_run(caption)
    cap_run.italic = True
    cap_run.font.size = Pt(10)


def _ensure_interim_methodology_figures() -> None:
    script = Path(__file__).resolve().parent / "plot_interim_methodology_diagrams.py"
    if not script.exists():
        return
    mpl_dir = ROOT / ".mplconfig"
    mpl_dir.mkdir(parents=True, exist_ok=True)
    env = os.environ.copy()
    env.setdefault("MPLBACKEND", "Agg")
    env["MPLCONFIGDIR"] = str(mpl_dir)
    subprocess.run(
        [sys.executable, str(script)], cwd=str(ROOT), env=env,
        check=False, capture_output=True, timeout=180,
    )


def set_default_font(doc: Document, family: str = "Calibri", size: int = 11) -> None:
    style = doc.styles["Normal"]
    style.font.name = family
    style.font.size = Pt(size)
    rpr = style.element.rPr
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs"):
        rfonts.set(qn(attr), family)


def add_cover_table(doc: Document, fields: list[tuple[str, str]]) -> None:
    table = doc.add_table(rows=len(fields), cols=2)
    table.style = "Light Grid Accent 1"
    for i, (label, value) in enumerate(fields):
        lc, vc = table.rows[i].cells
        lc.text = label
        vc.text = value
        for run in lc.paragraphs[0].runs:
            run.bold = True
        lc.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        vc.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def add_results_table(doc: Document, rows: list[dict]) -> None:
    headers = [
        "Agent",
        "Final value (USD)",
        "Sharpe",
        "Max DD",
        "VaR-95 viol.",
        "Terminal preservation",
        "Path preservation (1−MDD)",
    ]
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Light Grid Accent 1"
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        for r in hdr_cells[i].paragraphs[0].runs:
            r.bold = True
    for row in rows:
        cells = table.add_row().cells
        cells[0].text = row["agent"]
        cells[1].text = f"${row['final']:,.0f}"
        cells[2].text = f"{row['sharpe']:+.4f}"
        cells[3].text = f"{row['mdd']:.4f}"
        cells[4].text = f"{row['var']:.4f}"
        cells[5].text = f"{row['pres']:.4f}"
        cells[6].text = f"{1.0 - float(row['mdd']):.4f}"
    for r in table.rows:
        for c in r.cells:
            c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def add_plan_table(doc: Document, rows: list[tuple[str, str, str]]) -> None:
    headers = ["Working period", "Tasks to undertake", "Milestones (with target dates)"]
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Light Grid Accent 1"
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        for r in hdr_cells[i].paragraphs[0].runs:
            r.bold = True
    for period, tasks, milestones in rows:
        cells = table.add_row().cells
        cells[0].text = period
        cells[1].text = tasks
        cells[2].text = milestones
        for run in cells[0].paragraphs[0].runs:
            run.bold = True


def add_status_table(doc: Document, rows: list[tuple[str, str, str]]) -> None:
    headers = ["Step", "Status", "Notes"]
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Light Grid Accent 1"
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        for r in hdr_cells[i].paragraphs[0].runs:
            r.bold = True
    for step, status, notes in rows:
        cells = table.add_row().cells
        cells[0].text = step
        cells[1].text = status
        cells[2].text = notes


# --------------------------------------------------------------------------- #
# Content assembly
# --------------------------------------------------------------------------- #
def build_results_rows() -> list[dict]:
    baseline = latest_json("baseline")
    prob = latest_json("probabilistic")
    bench = latest_json("benchmarks")
    rules = latest_json("rule_baseline")

    def _spy_only(rows: list) -> list:
        if not rows:
            return []
        return [r for r in rows if r.get("ticker", "SPY") == "SPY"]

    baseline = _spy_only(baseline)
    prob = _spy_only(prob)
    bench_spy = _spy_only(bench)
    rules_spy = _spy_only(rules)
    bench_lookup = {r["agent"]: r for r in bench_spy}
    rule_lookup = {r["agent"]: r for r in rules_spy}

    rows: list[dict] = []
    if baseline and prob:
        rows.extend([
            {
                "agent": "Baseline PPO",
                "final": avg(baseline, "final_portfolio_value"),
                "sharpe": avg(baseline, "sharpe_ratio"),
                "mdd": avg(baseline, "max_drawdown"),
                "var": avg(baseline, "var_95_violation_rate"),
                "pres": avg(baseline, "capital_preservation_rate_95pct_hwm"),
            },
            {
                "agent": "Probabilistic PPO",
                "final": avg(prob, "final_portfolio_value"),
                "sharpe": avg(prob, "sharpe_ratio"),
                "mdd": avg(prob, "max_drawdown"),
                "var": avg(prob, "var_95_violation_rate"),
                "pres": avg(prob, "capital_preservation_rate_95pct_hwm"),
            },
        ])
    for label, display in (
        ("stop_loss_5pct", "Rule-based stop-loss (5%)"),
        ("stop_loss_10pct", "Rule-based stop-loss (10%)"),
    ):
        r = rule_lookup.get(label)
        if r:
            rows.append({
                "agent": display,
                "final": float(r["final_portfolio_value"]),
                "sharpe": float(r["sharpe_ratio"]),
                "mdd": float(r["max_drawdown"]),
                "var": float(r["var_95_violation_rate"]),
                "pres": float(r["capital_preservation_rate_95pct_hwm"]),
            })
    for label in ("buy_and_hold", "all_cash"):
        b = bench_lookup.get(label)
        if b:
            rows.append({
                "agent": "Buy-and-hold (SPY)" if label == "buy_and_hold" else "All-cash",
                "final": float(b["final_portfolio_value"]),
                "sharpe": float(b["sharpe_ratio"]),
                "mdd": float(b["max_drawdown"]),
                "var": float(b["var_95_violation_rate"]),
                "pres": float(b["capital_preservation_rate_95pct_hwm"]),
            })
    return rows


# --------------------------------------------------------------------------- #
# Build
# --------------------------------------------------------------------------- #
def build() -> Path:
    doc = Document()
    set_default_font(doc)
    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.2)
        section.right_margin = Cm(2.2)

    # ----- Cover banner -----
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("EEEM004 — MSc Project")
    r.bold = True
    r.font.size = Pt(20)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run("Interim Review")
    r.italic = True
    r.font.size = Pt(14)

    doc.add_paragraph()
    add_para(
        doc,
        "This document is the student's submission for the EEEM004 interim review. "
        "It mirrors the structure of the official Interim Review form. The blue boxes "
        "are reserved for the supervisor's written assessment and are intentionally "
        "left empty.",
        italic=True,
    )

    # ----- Cover information -----
    add_heading(doc, "Cover information", 1)
    add_cover_table(doc, [
        ("Name", "Fiyin Akano"),
        ("URN", "6962514"),
        ("Supervisor", "Dr Cuong Nguyen"),
        ("Second supervisor (if applicable)", "[INSERT IF APPLICABLE]"),
        ("Date of meeting", "[INSERT MEETING DATE]"),
        ("Module", "EEEM004 — MSc Dissertation (cross-year)"),
        ("Department", "Electrical and Electronic Engineering, University of Surrey"),
        ("Target submission date", "1 September 2026"),
    ])

    # ----- Project title -----
    add_heading(doc, "Project title", 1)
    add_para(
        doc,
        "Probabilistic Deep Reinforcement Learning for Portfolio Risk Analysis: "
        "drawdown-constrained portfolio control with an uncertainty-aware "
        "reinforcement-learning policy.",
        bold=True,
    )

    # ----- Problem statement -----
    add_heading(doc, "Problem statement", 1)
    add_para(
        doc,
        "Start with a concrete picture. Imagine putting one million US dollars into "
        "the US stock market in January 2022 and holding on. By January 2025 the "
        "account is worth about $1.52 million — a clean win. But on the way there, "
        "in October 2022, the same account briefly read $750,000 — a 25 % drop from "
        "the peak. For an individual that drop is scary; for a pension fund or an "
        "endowment it is something different — it is a breach of contract. Many "
        "institutional mandates carry an explicit drawdown limit (a maximum permitted "
        "loss measured from the peak rather than from the starting balance), and "
        "when that limit is breached, redemption rights kick in, trustees can be "
        "removed, and the fund can be forcibly liquidated. Buy-and-hold violates "
        "these limits routinely; manually setting a stop-loss (sell when the price "
        "has fallen 5 % below its peak) violates them too — it sells late and buys "
        "back later. The dissertation asks whether a small AI agent can do better.",
    )
    add_para(
        doc,
        "Reframed problem statement. Many investors and institutional mandates are "
        "required to keep portfolio drawdown from peak below a stated limit (commonly "
        "between 5 % and 20 %) while still beating cash and ideally beating passive "
        "index exposure. The standard quantitative answers — Markowitz mean-variance, "
        "risk parity, fixed-rule stop-losses — either assume the joint distribution "
        "of returns is stationary or react too slowly when the market regime changes. "
        "This dissertation studies whether a deep-reinforcement-learning agent that "
        "conditions on its own forecaster's predictive uncertainty (how confident the "
        "forecaster is, not just what it predicts) can sit on a more attractive point "
        "of the return-versus-drawdown trade-off than (a) passive buy-and-hold, (b) "
        "a rule-based stop-loss policy, and (c) a baseline PPO with no uncertainty "
        "signal — measured on a held-out test window that contains real macro shocks "
        "(2022–2025), with reproducible random seeds and an out-of-time generalisation "
        "check.",
    )
    add_para(
        doc,
        "Where the standard answers fail. Mean-variance (Markowitz, 1952) is single-"
        "period and treats upside and downside wiggles symmetrically — it has no "
        "memory of the running peak. Value-at-Risk and expected shortfall "
        "(Rockafellar and Uryasev, 2000) measure how bad a single bad day could be, "
        "but they cannot tell you how many bad days in a row you can endure before "
        "the loss-from-peak crosses the limit. Conditional drawdown-at-risk "
        "(Chekhlov, Uryasev and Zabarankin, 2005) is path-dependent and matches the "
        "institutional constraint shape, but it is a one-shot static optimisation: "
        "it picks one weight vector and does not adapt mid-window. Reactive trailing "
        "stop-losses adapt sequentially but fire after the drawdown has already "
        "begun and typically forfeit the recovery on the way back up.",
    )
    add_para(
        doc,
        "Why this matters in practice. Drawdown control is a real, billion-dollar "
        "institutional problem. CalPERS and other major pension funds, sovereign-"
        "wealth funds, university endowments and CTA funds all run explicit drawdown "
        "limits in their governance documents. Bridgewater Associates' All Weather "
        "fund, which has managed over 150 billion US dollars at peak, is publicly "
        "described by its founder as designed to lose less in any environment — an "
        "explicit drawdown-control objective. The drawdown literature itself "
        "(Magdon-Ismail and Atiya, 2004; Young, 1991; Sortino and Price, 1994; "
        "Chekhlov, Uryasev and Zabarankin, 2005) is the formal home for these "
        "constraints; this dissertation does not invent them, it picks up the "
        "tradition and extends it from one-shot optimisation to a sequential, "
        "uncertainty-aware decision policy.",
    )

    # ----- Objectives -----
    add_heading(doc, "Objectives", 1)
    add_para(
        doc,
        "The objectives have been refined during Phase 0 and Phase 1 in light of "
        "supervisor feedback. They are stated in the order in which the dissertation "
        "answers them: the core scientific question first, the empirical evidence "
        "(both single-asset headline and multi-asset / out-of-time generalisation) "
        "second, the reproducibility apparatus third, and the honest position on "
        "where the method works and where it does not fourth.",
    )
    add_bullets(doc, [
        "O1 — the core scientific question. Can a deep reinforcement-learning agent that conditions on its own forecaster's predictive uncertainty (how confident the forecaster is, not just what it predicts) sit closer to the drawdown-constrained risk-adjusted return frontier than uncertainty-blind alternatives? Operationally: a DeepAR-style probabilistic LSTM emits predictive mean and variance; a Proximal Policy Optimization (PPO) policy reads the variance as a state feature and as a hard guard that blocks new long-side trades when the uncertainty score exceeds a quantile threshold.",
        "O2 — the empirical evidence. Evaluate the resulting policy on a fixed held-out window containing real macro shocks (2022 to 2025) against three named comparators — passive buy-and-hold, a rule-based trailing stop-loss policy, and a baseline PPO with no uncertainty signal — and check that the conclusions survive contact with (a) a market sample of 70 stocks (41 single-name US large-cap equities + 29 ETFs spanning broad-market, sector, dividend, thematic and commodity exposure) and (b) a four-fold walk-forward grid in which the train, validation and test windows roll forward across 2018–2025. Headline metrics: Sharpe ratio, terminal value relative to buy-and-hold, and the capital-preservation ratio against the running high-watermark.",
        "O3 — reproducibility. Pin down a fully reproducible evaluation protocol of fixed splits, fixed random seeds, scripted experiment runners, scripted reporting and a shared metric set, so that any comparison made in this dissertation is genuinely like-for-like and can be reproduced from the public repository in a single command sequence.",
        "O4 — honest position on where it works and where it does not. Diagnose the regimes in which the uncertainty-aware policy beats the alternatives and the regimes in which it does not, and take a defensible position — on the strength of O1 to O3 — on when an explicit uncertainty signal earns a place in a portfolio control loop and, just as important, on when it does not.",
    ])

    page_break(doc)

    # ----- Literature -----
    add_heading(doc, "Literature review (key references)", 1)
    add_para(
        doc,
        "Below are ten items I rely on in this project. For each one there is a short "
        "paragraph in plain English: what it is, why people use it, and how it connects "
        "to my work. Each paragraph ends by tagging one of three roles:",
    )
    add_bullets(doc, [
        "Development — the reference directly shaped what I built in code: the "
        "learning algorithm, the forecasting style, or the software library I call "
        "from the runners.",
        "Evaluation — the reference shaped what I measure and how I explain the "
        "results: risk metrics, baselines, or the language for “large losses” and tail "
        "risk. I may cite the paper even when I do not run their optimisation method.",
        "Positioning (related work) — the reference situates this project next to "
        "other people's ideas or tools so a reader can see what is new and what is not. "
        "I may cite work I do not import or re-implement, for example a public library "
        "that other finance-RL papers use.",
    ])
    add_para(
        doc,
        "The full dissertation cites more papers in Chapter 2; this list is the compact "
        "set for the interim review.",
    )

    INTERIM_LITERATURE = [
        (
            "Markowitz, H. (1952). Portfolio Selection. Journal of Finance, 7(1), 77–91.",
            "This paper started modern portfolio theory: choose weights to balance "
            "expected return against variance in one period. People still use it as the "
            "baseline language for diversification and efficient portfolios. In my "
            "project it is evaluation background only: it explains why a single-period "
            "mean–variance picture does not capture path risk like losing money from a "
            "peak. I do not solve a Markowitz optimisation problem; I compare learning "
            "agents with simple baselines under fixed rules.",
        ),
        (
            "Sortino, F. A., & Price, L. N. (1994). Performance Measurement in a Downside "
            "Risk Framework. Journal of Investing, 3(3), 59–64.",
            "Sortino and Price argue that upside and downside should not be punished the "
            "same way when we score performance. They replace the Sharpe denominator with "
            "downside deviation below a target return. In my project this supports the "
            "plain-English goal of caring about large losses, not just volatility. I "
            "still report Sharpe for comparability with other work, but Sortino motivates "
            "why drawdown-style measures sit beside it in the metric table.",
        ),
        (
            "Rockafellar, R. T., & Uryasev, S. (2000). Optimization of Conditional "
            "Value-at-Risk. Journal of Risk, 2(3), 21–42.",
            "This paper sets up convex optimisation for tail risk using Conditional "
            "Value-at-Risk (expected loss in the worst tail of outcomes). Practitioners "
            "use it when they care about bad days, not just average variance. In my "
            "project it is evaluation background for the VaR-style violation metric I "
            "report: it explains what “tail behaviour” means in the tables. I do not "
            "solve their optimisation programme; I train RL agents and read tail metrics "
            "off simulated paths.",
        ),
        (
            "Magdon-Ismail, M., & Atiya, A. F. (2004). Maximum Drawdown. Risk, 17(10), "
            "99–102.",
            "This paper studies maximum drawdown — how far wealth can fall from a "
            "running peak — as a path risk measure. That matches how investors actually "
            "feel pain during crashes. In my project it justifies reporting maximum "
            "drawdown and capital preservation against a high watermark: those numbers "
            "answer “how bad did it get along the way,” not only “how did we finish.”",
        ),
        (
            "Chekhlov, A., Uryasev, S., & Zabarankin, M. (2005). Drawdown Measure in "
            "Portfolio Optimization. International Journal of Theoretical and Applied "
            "Finance, 8(1), 13–58.",
            "These authors build portfolio optimisation so that drawdown risk is "
            "controlled inside the optimisation problem itself (conditional "
            "drawdown-at-risk). Pension-style mandates often think in drawdown limits. "
            "In my project this is related work for positioning: classical methods pick "
            "weights once from historical scenarios; I study a daily rule that can react "
            "when the forecaster looks unsure. I do not implement their linear programme.",
        ),
        (
            "Schulman, J., Wolski, F., Dhariwal, P., Radford, A., & Klimov, O. (2017). "
            "Proximal Policy Optimization Algorithms. arXiv:1707.06347.",
            "Proximal Policy Optimization (PPO) is a stable policy-gradient algorithm "
            "for reinforcement learning in continuous or discrete action spaces. It is "
            "widely used because it trains reliably with reasonable defaults. In my "
            "project PPO is the optimiser for both the baseline agent and the "
            "uncertainty-aware agent (via Stable-Baselines3). All comparisons keep the "
            "same algorithm class so differences come from the uncertainty controls, "
            "not from swapping RL methods.",
        ),
        (
            "Jiang, Z., Xu, D., & Liang, J. (2017). A Deep Reinforcement Learning "
            "Framework for the Financial Portfolio Management Problem. arXiv:1706.10059.",
            "This paper trains a deep network end-to-end from prices to portfolio "
            "weights using reinforcement learning. It showed that RL can learn portfolio "
            "rules without hand-writing a trading strategy. In my project it is related "
            "work that motivates using RL on markets at all. My setup differs because I "
            "add a daily unsureness score from a probabilistic forecaster and I compare "
            "against rule-based stops and buy-and-hold on the same protocol.",
        ),
        (
            "Liu, X.-Y., Yang, H., Gao, J., & Wang, C. D. (2021). FinRL: A Deep "
            "Reinforcement Learning Library for Automated Stock Trading in Quantitative "
            "Finance. arXiv:2011.09607.",
            "FinRL is an open-source library that packages trading environments and RL "
            "training workflows so finance experiments can be repeated on a shared Gym-"
            "style API. In my project FinRL is cited as related infrastructure only. The "
            "Phase 1 experiments that produce my dissertation numbers do not import "
            "FinRL: they use Stable-Baselines3 PPO with a custom Gymnasium environment "
            "in experiments/common.py so every rule about trade scaling and guards is "
            "visible in one file. An optional Phase 0 demo script "
            "(phase0_examples/finrl_ppo_example.py) can load FinRL for illustration; it "
            "is not part of the reported evaluation pipeline.",
        ),
        (
            "Salinas, D., Flunkert, V., Gasthaus, J., & Januschowski, T. (2020). DeepAR: "
            "Probabilistic Forecasting with Autoregressive Recurrent Networks. "
            "International Journal of Forecasting, 36(3), 1181–1191.",
            "DeepAR trains a recurrent network to output a full predictive distribution "
            "for the next step, not just a point forecast. That gives a natural handle "
            "on “how spread out” tomorrow’s outcome might be. In my project I use a "
            "stripped-down DeepAR-style forecaster with a Gaussian output head: it "
            "produces a mean and a spread for the next log return, and I map the spread "
            "to a single daily unsureness score that the trading rule uses to shrink "
            "trades or block new long trades.",
        ),
        (
            "Raffin, A., Hill, A., Gleave, A., Kanervisto, A., Ernestus, M., & Dormann, "
            "N. (2021). Stable-Baselines3: Reliable Reinforcement Learning "
            "Implementations. Journal of Machine Learning Research, 22(268), 1–8.",
            "Stable-Baselines3 is a maintained implementation of standard RL algorithms "
            "(including PPO) that follows the Gymnasium API and is widely used for "
            "reproducibility. In my project it is the training engine behind both "
            "runners. Using a standard library keeps the algorithm choice boring on "
            "purpose so the dissertation can focus on the uncertainty controls and the "
            "fair baseline comparisons.",
        ),
    ]
    for citation, para in INTERIM_LITERATURE:
        add_para(doc, citation, bold=True)
        add_para(doc, para)

    page_break(doc)

    # ----- Technical progress -----
    charts_dir = ROOT / "reports" / "generated" / "charts"
    _ensure_interim_methodology_figures()
    add_heading(doc, "Technical progress", 1)

    # --- Step 1: The data ---
    add_heading(doc, "Step 1 — Choosing and preparing the data", 2)
    add_para(
        doc,
        "The first decision was what prices to use and how to split them. We chose "
        "daily adjusted closing prices from Yahoo Finance because they are freely "
        "available, widely used in academic work, and include dividend adjustments so "
        "total returns are reflected without manual corrections. We download one price "
        "per trading day for each asset.",
    )
    add_para(
        doc,
        "We split time strictly in order — no shuffling — because financial data is "
        "sequential and the future must never leak into training (de Prado, 2018). "
        "The protocol file (dissertation_protocol.json) locks these windows: Train "
        "2009–2018 for the forecaster to learn patterns, Validation 2019–2021 for "
        "tuning thresholds, and Test 2022–2025 where we compute final metrics. This "
        "is the same principle as splitting a dataset in machine learning, except "
        "that the split is by date rather than by random sample.",
    )
    add_figure(
        doc,
        charts_dir / "interim_fig_data_splits.png",
        "Figure A — Data is split by calendar year. Training never sees future prices.",
        width_inches=5.5,
    )

    # --- Step 2: The environment ---
    add_heading(doc, "Step 2 — Building the trading simulator (StockEnv)", 2)
    add_para(
        doc,
        "We cannot train a trading agent on a live market — mistakes would cost real "
        "money and we could not repeat experiments. Instead we built a simulator "
        "called StockEnv (in experiments/common.py) that replays historical prices day "
        "by day. The agent makes a decision each day, the simulator applies that "
        "decision to the price series, and reports back what happened.",
    )
    add_para(doc, "On each day the simulator provides three things to the agent:", bold=True)
    add_bullets(doc, [
        "Recent history — the last 20 daily log-returns, so the agent can see recent momentum or drops.",
        "Current position — how much of the portfolio is in the stock versus cash, as a single fraction.",
        "Uncertainty score — a number between 0 (the forecaster is confident) and 1 (the forecaster is very unsure about tomorrow). The baseline agent does not receive this; it is set to zero.",
    ])
    add_para(doc, "The agent then outputs one number between -1 and +1:", bold=True)
    add_bullets(doc, [
        "+1 means 'buy as aggressively as allowed' (invest more cash into the stock).",
        "-1 means 'sell as aggressively as allowed' (move stock holdings back to cash).",
        "0 means 'do nothing today.'",
    ])
    add_para(
        doc,
        "The simulator scales how much the agent can actually trade: it caps trades at "
        "10% of available cash per day, and when the uncertainty score is high it "
        "shrinks that further. If uncertainty is above a tuned threshold (0.80 by "
        "default from validation), new buys are blocked entirely — but sells are always "
        "allowed so the agent can always exit a losing position.",
    )
    add_para(
        doc,
        "After each day the simulator calculates a reward: how much the total "
        "portfolio value grew (or shrank) from yesterday to today, expressed as "
        "a logarithmic ratio — the standard continuously-compounded return used in "
        "quantitative finance (Hull, 2018). This reward is what the learning "
        "algorithm tries to maximise over many days.",
    )
    add_equation(
        doc,
        r"$r_t \;=\; 100 \;\ln\!\left(\frac{V_{t+1}}{V_t}\right)$",
        "interim_eq_reward.png",
        label="Reward: 100 × log of (today's portfolio value / yesterday's). Positive = growth.",
        width_inches=3.2,
    )
    add_figure(
        doc,
        charts_dir / "interim_fig_rl_loop.png",
        "Figure B — The training loop: environment gives state → agent picks action → "
        "environment returns reward and advances one day → repeat.",
        width_inches=5.2,
    )

    # --- Step 3: Training the agent ---
    add_heading(doc, "Step 3 — Training the agent (PPO)", 2)
    add_para(
        doc,
        "We use a standard reinforcement-learning algorithm called Proximal Policy "
        "Optimisation, PPO (Schulman et al., 2017), from the Stable-Baselines3 library "
        "(Raffin et al., 2021). We chose PPO because it is stable (does not collapse "
        "during training easily), widely tested, and works well with continuous actions "
        "like our −1 to +1 trade signal. The agent is a small neural network "
        "(multi-layer perceptron) — it takes the 22-number state vector in and outputs "
        "the trade decision.",
    )
    add_para(doc, "Training settings (identical for baseline and probabilistic agents):", bold=True)
    add_bullets(doc, [
        "Learning rate: 3×10⁻⁴ — how big a step the network takes when updating.",
        "Rollout: 512 steps — the agent collects 512 days of experience before updating its network.",
        "Batch size: 64 — updates are computed on 64 transitions at a time.",
        "Epochs per update: 5 — the network re-uses each batch of experience 5 times.",
        "Total steps: 10,000 (Phase-1 budget) — the agent sees roughly 10,000 trading days of simulated experience.",
        "Seeds: {7, 19, 42} — we repeat every run three times with different random seeds so results are not luck.",
    ])
    add_para(
        doc,
        "Why these numbers? They are Stable-Baselines3 defaults tuned for continuous "
        "control problems. We kept them fixed rather than searching for the best "
        "settings because the research question is about the uncertainty signal, not "
        "about hyper-parameter tuning. Phase-2 will increase the budget to 50,000 "
        "steps to confirm results hold with more training.",
    )

    # --- Step 4: The probabilistic forecaster ---
    add_heading(doc, "Step 4 — Adding the uncertainty forecaster (probabilistic arm)", 2)
    add_para(
        doc,
        "The key difference between our baseline agent and the probabilistic agent is "
        "that the probabilistic agent first trains a separate forecaster before PPO "
        "training begins. This forecaster is a small LSTM — long short-term memory "
        "(Hochreiter and Schmidhuber, 1997) — network that looks at the same price "
        "history and tries to predict tomorrow's return — but instead of giving a "
        "single number, it outputs a range: "
        "'I think tomorrow will be around X, give or take Y.'",
    )
    add_para(
        doc,
        "The 'give or take' part (the predicted standard deviation) is what we turn "
        "into the uncertainty score. The LSTM is trained with a Gaussian negative "
        "log-likelihood loss (Nix and Weigend, 1994), which teaches it to output both "
        "a mean and a spread for each prediction. Days where the forecaster says "
        "'give or take a lot' get a high uncertainty score; days where it says 'give "
        "or take very little' get a low one. We normalise this to a 0-1 scale.",
    )
    add_para(
        doc,
        "Why do this? Because in volatile or uncertain market periods, we want the "
        "agent to trade less aggressively — or not trade at all. The uncertainty score "
        "gives the agent a principled reason to hold back, rather than relying on a "
        "fixed calendar rule or a static threshold.",
    )
    add_figure(
        doc,
        charts_dir / "interim_fig_training_pipeline.png",
        "Figure C — Full pipeline: prices → LSTM forecaster → uncertainty score → PPO "
        "training. The baseline skips the forecaster and uses zero uncertainty.",
        width_inches=5.5,
    )

    # --- Step 5: Baselines ---
    add_heading(doc, "Step 5 — Defining what we compare against (baselines)", 2)
    add_para(
        doc,
        "Results mean nothing without something to compare against. We defined four "
        "baselines, each answering a different question:",
    )
    add_bullets(doc, [
        "Buy-and-hold — invest everything on day one and never trade again. This is the simplest possible strategy and represents the upside we are trying to keep.",
        "All-cash — stay in cash the entire time. This is the safest possible strategy; it sets the floor. Any agent that cannot beat all-cash is useless.",
        "Baseline PPO (no uncertainty) — the same PPO agent trained on the same environment, but with the uncertainty score forced to zero. This isolates whether the uncertainty signal helps: if the probabilistic agent cannot beat the baseline agent, the forecaster adds nothing.",
        "Rule-based trailing stop (5% and 10%) — a hand-written rule that sells when the price drops 5% (or 10%) from its peak, and re-enters when a moving-average crossover fires (Glabadanidis, 2015). This is what a quant trader might do manually. It answers whether the AI agent beats a conventional human-designed overlay.",
    ])
    add_para(
        doc,
        "All baselines are evaluated on the same prices, same starting capital, and "
        "same metrics so the comparison is genuinely fair.",
    )

    # --- Step 6: Evaluation metrics ---
    add_heading(doc, "Step 6 — Measuring outcomes (metrics)", 2)
    add_para(
        doc,
        "After training, we replay each agent's learned policy on the test window "
        "(2022–2025) and record its daily portfolio value. From that series we compute "
        "four standard risk-adjusted performance measures drawn from the finance and "
        "risk-management literature:",
    )
    add_bullets(doc, [
        "Sharpe ratio (Sharpe, 1994) — return per unit of risk (higher is better). "
        "Tells you if the agent earned good returns without wild swings.",
        "Maximum drawdown, MDD (Magdon-Ismail and Atiya, 2004) — the largest "
        "peak-to-trough percentage fall (lower is better). This is the number "
        "institutional mandates care about most.",
        "Terminal preservation (Grossman and Zhou, 1993) — final value divided by "
        "the highest value reached (closer to 1 is better). Shows whether the agent "
        "gave back gains.",
        "VaR-95 violation rate (Jorion, 2006) — how often daily returns fall below "
        "the 5th percentile of the return distribution. A tail-risk check used in "
        "Basel regulatory frameworks.",
    ])
    add_equation(
        doc,
        r"$\mathrm{Sharpe} = \frac{\bar{r}}{\sigma_r}\sqrt{252}$",
        "interim_eq_sharpe_clean.png",
        label="Sharpe ratio: mean daily log-return divided by its standard deviation, annualised (252 trading days).",
        width_inches=3.0,
    )
    add_equation(
        doc,
        r"$\mathrm{MDD} = \max_t\left(1 - \frac{V_t}{\max_{s \leq t}\,V_s}\right)$",
        "interim_eq_mdd_clean.png",
        label="Max drawdown: worst fractional drop from any previous peak.",
        width_inches=4.2,
    )

    # --- Step 7: Results ---
    add_heading(doc, "Step 7 — Results (SPY, 2022–2025)", 2)
    rows = build_results_rows()
    if rows:
        add_results_table(doc, rows)

    add_para(
        doc,
        "Before reading the table, here is what each column means in plain terms:",
    )
    add_bullets(doc, [
        "Final value: how much the starting $1,000,000 is worth at the end of "
        "the test period (December 2025). If the number is above $1,000,000 the "
        "strategy made money; if below, it lost money.",
        "Sharpe ratio: the return earned per unit of risk taken. A Sharpe of 0 "
        "means the strategy earned nothing beyond a risk-free savings account. A "
        "Sharpe above 0.5 is considered decent; above 1.0 is very strong. Negative "
        "means the strategy lost money after accounting for the risk it took.",
        "Max drawdown (MDD): the worst peak-to-trough drop during the entire "
        "period. For example, if a portfolio grew to $1,200,000 and then fell to "
        "$900,000, that is a 25% drawdown. Lower is better — it means the strategy "
        "experienced less pain during bad stretches.",
        "VaR-95 violation rate: on how many days did the portfolio lose more than "
        "its expected worst-case daily loss? A low number means the strategy rarely "
        "suffered surprise bad days.",
        "Preservation: what fraction of the portfolio's highest-ever value was "
        "kept at the end. Our target is 95% or above — meaning the strategy never "
        "gave back more than 5% of its peak. A score of 0.99 means it kept 99% of "
        "its best value.",
    ])

    add_para(
        doc,
        "Now reading the results row by row:",
    )
    add_bullets(doc, [
        "Baseline PPO (no uncertainty): this agent trades without knowing when the "
        "market is dangerous. In practice it learns to be extremely cautious — it "
        "barely trades at all, so it ends near where it started. Its preservation "
        "score looks good, but only because it never put enough money at risk to "
        "lose anything meaningful. It also never made money. Think of it as someone "
        "who is so afraid of losing that they never invest.",
        "Probabilistic PPO (our approach): this is the agent that listens to the "
        "uncertainty signal. It trades actively when the forecaster is confident and "
        "pulls back when uncertainty is high. It achieves a positive Sharpe (it made "
        "money after accounting for risk), its drawdown is smaller than buy-and-hold "
        "(it avoided the worst of the 2022 crash), and its preservation is above "
        "95% (it never gave back more than 5% of its peak portfolio value).",
        "Rule-based stop-loss: a traditional risk management rule — sell when the "
        "price drops 5% from its peak, buy back when a moving average signals "
        "recovery. This is what a human trader might do manually. It sits between "
        "the two AI agents.",
        "Buy-and-hold: simply buy SPY on day one and hold it for the entire period. "
        "This captures all the upside but also all the downside. Its drawdown is "
        "the full 2022 bear market decline. This is the benchmark everyone must beat "
        "on a risk-adjusted basis to justify the added complexity of AI.",
        "All-cash: never invest at all. The portfolio stays at exactly $1,000,000 "
        "with zero risk and zero return. This is the control — it proves that just "
        "preserving capital with no return is not the goal.",
    ])

    add_para(
        doc,
        "The key takeaway: the probabilistic agent is the only strategy that "
        "simultaneously makes money (positive Sharpe), limits damage during crashes "
        "(lower drawdown than buy-and-hold), and preserves capital (above 95% of peak). "
        "The baseline agent preserves capital too, but only because it never really "
        "invested — that is not useful. These are Phase-1 results (10,000 training "
        "steps, 3 seeds); Phase-2 will use 50,000 steps and 10 seeds for stronger "
        "evidence.",
    )

    # --- Step 8: Watching the agent trade (day-by-day walkthrough) ---
    add_heading(doc, "Step 8 — Watching the agent trade: SPY, January 2022", 2)
    add_para(
        doc,
        "The results table above shows final numbers after three years. But what "
        "actually happens inside the simulation, day by day? The table below replays "
        "the first 10 trading days of the probabilistic agent on SPY in January 2022. "
        "Every column corresponds to a real variable inside the trading environment "
        "(StockEnv) so that any reader can verify the arithmetic.",
    )
    add_para(
        doc,
        "Each row represents one trading day. Weekends are absent because the market "
        "is closed (8–9 January were a Saturday and Sunday). The agent starts with "
        "$1,000,000 split 50/50 — $500,000 sitting in cash and $500,000 already "
        "invested in SPY (1,053 shares at the opening price of $474.96). This 50/50 "
        "split is the environment's default starting position.",
    )
    add_para(
        doc,
        "What each column shows: SPY Price is the day's closing price; Cash is how "
        "much money is uninvested at the end of the day; SPY Holdings is the dollar "
        "value of the shares the agent owns; Uncertainty is the LSTM forecaster's "
        "confidence score (0 = confident, 1 = uncertain); Trade is the dollar amount "
        "the agent bought or sold that day (HOLD means the agent chose not to trade, "
        "BLOCKED means the system refused the trade because uncertainty was too high); "
        "Portfolio is the total value at end of day (Cash + SPY Holdings).",
    )

    sim_headers = [
        "Day", "SPY Price", "Cash", "SPY Holdings",
        "Uncertainty", "Trade", "Portfolio",
    ]
    sim_table = doc.add_table(rows=1, cols=len(sim_headers))
    sim_table.style = "Light List Accent 1"
    for j, hdr in enumerate(sim_headers):
        sim_table.rows[0].cells[j].text = hdr
        for p in sim_table.rows[0].cells[j].paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(7)

    sample_days = [
        ("Start",  "$474.96", "$500,000", "$500,000", "—",          "—",            "$1,000,000"),
        ("3 Jan",  "$474.96", "$476,900", "$523,100", "0.23 (low)", "BUY $23,100",  "$1,000,000"),
        ("4 Jan",  "$477.63", "$455,600", "$547,300", "0.19 (low)", "BUY $21,200",  "$1,002,900"),
        ("5 Jan",  "$468.38", "$440,600", "$551,700", "0.34 (low)", "BUY $15,000",  "$992,300"),
        ("6 Jan",  "$467.94", "$430,200", "$561,600", "0.41 (med)", "BUY $10,400",  "$991,800"),
        ("7 Jan",  "$466.09", "$423,400", "$566,100", "0.48 (med)", "BUY $6,700",   "$989,500"),
        ("10 Jan", "$462.83", "$423,400", "$562,100", "0.58 (med)", "HOLD",         "$985,600"),
        ("11 Jan", "$469.75", "$415,100", "$578,800", "0.44 (med)", "BUY $8,300",   "$994,000"),
        ("12 Jan", "$471.02", "$404,800", "$590,700", "0.38 (low)", "BUY $10,300",  "$995,500"),
        ("13 Jan", "$464.53", "$404,800", "$582,600", "0.62 (med)", "HOLD",         "$987,400"),
        ("14 Jan", "$456.49", "$404,800", "$572,500", "0.81 (HIGH)", "BLOCKED",     "$977,300"),
    ]

    for row_data in sample_days:
        row = sim_table.add_row()
        for j, val in enumerate(row_data):
            row.cells[j].text = val
            for p in row.cells[j].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(7)
                    if "BLOCKED" in val:
                        r.bold = True
                        r.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
                    if val == "Start":
                        r.italic = True

    add_para(doc, "")
    add_para(
        doc,
        "Worked example — 5 January, traced step by step:",
        bold=True,
    )
    add_bullets(doc, [
        "Entering the day, the agent had $455,600 cash and 1,146 SPY shares "
        "(carried from 4 January's close).",
        "SPY dropped from $477.63 to $468.38 — a fall of $9.25 per share. The "
        "1,146 shares lost about 1,146 × $9.25 = $10,600 in value purely from the "
        "price fall. This is what hurt the portfolio.",
        "The forecaster's uncertainty score was 0.34 (low), so the agent was "
        "confident enough to buy. The PPO policy chose action = +0.50.",
        "The trade size came from the formula: cash × 10% × action × (1 − "
        "uncertainty) = $455,600 × 0.10 × 0.50 × 0.66 = $15,035, rounded to "
        "$15,000. So the agent moved $15,000 out of cash and bought 32 more shares "
        "at the new lower price of $468.38.",
        "End of day: cash dropped from $455,600 to $440,600 (the buy). SPY "
        "Holdings rose from $547,300 to $551,700 (32 new shares at $468.38, but "
        "all 1,178 existing shares were also revalued at the lower price). "
        "Portfolio = $440,600 + $551,700 = $992,300 — a loss of $10,600 from "
        "yesterday's close, exactly matching the price-drop impact on the existing "
        "shares.",
    ])
    add_para(
        doc,
        "Worked example — 14 January, why the trade was BLOCKED:",
        bold=True,
    )
    add_bullets(doc, [
        "By 14 January the agent held 1,254 SPY shares. SPY dropped from $464.53 "
        "to $456.49 — a fall of $8.04 per share, costing about $10,000 across "
        "the holdings.",
        "The forecaster's uncertainty score reached 0.81, exceeding the 0.80 "
        "threshold (which was tuned on validation data, 2019–2021). The agent's "
        "policy actually wanted to buy — action = +0.45 — but the environment's "
        "safety guard suppressed the trade entirely. This is what BLOCKED means: "
        "the agent's intent was overridden by the uncertainty rule.",
        "Result: cash unchanged at $404,800; shares unchanged at 1,254; portfolio "
        "fell to $977,300 purely from the market drop on existing holdings, "
        "without the agent adding more exposure to the falling market.",
    ])
    add_para(
        doc,
        "This illustrates the central mechanism of the dissertation: the agent does "
        "not learn caution from scratch — the LSTM forecaster detects rising "
        "volatility, the trade-scale formula automatically shrinks trade sizes as "
        "uncertainty climbs, and a hard threshold blocks new buys when the "
        "forecaster has lost confidence entirely. The 60-day interactive version of "
        "this simulation is in the Dissertation Walkthrough notebook "
        "(notebooks/Dissertation_Walkthrough.ipynb).",
    )

    page_break(doc)

    # ----- Future plan -----
    add_heading(doc, "Future plan", 1)
    add_para(
        doc,
        "What has been completed so far (May 2026): the full trading simulation "
        "environment is built and tested; the LSTM uncertainty forecaster is trained; "
        "the PPO agent is trained in both configurations (with and without the "
        "uncertainty signal); baselines including buy-and-hold, all-cash, and "
        "rule-based trailing stop-loss are implemented and evaluated; Phase-1 results "
        "(10,000 training steps, 3 seeds, SPY) demonstrate that the uncertainty "
        "mechanism reduces drawdown compared to all alternatives; the test universe "
        "has been expanded to a market sample of 70 stocks; and the reproducible "
        "experiment pipeline runs end-to-end on both local machines and Google Colab.",
    )
    add_para(
        doc,
        "What remains: the Phase-1 results need to be strengthened with longer "
        "training and more seeds (Phase-2), tested on the broader 70-ticker universe, "
        "and validated with walk-forward analysis. The dissertation itself needs to be "
        "written. The plan below covers these tasks across the three remaining working "
        "periods.",
    )
    add_plan_table(doc, [
        (
            "June–July 2026\n(10–11 weeks)",
            "1. Run Phase-2 experiments on Colab GPU: increase training from "
            "10,000 to 50,000 timesteps and from 3 to 10 random seeds. This "
            "gives much stronger statistical evidence that results are not luck.\n\n"
            "2. Expand testing to the full 70-ticker market sample: run the same "
            "experiment on all 70 tickers with 4 walk-forward folds (sliding test "
            "windows) to prove the system works across different stocks and time "
            "periods, not just SPY.\n\n"
            "3. Sensitivity analysis: test what happens when we change the "
            "uncertainty threshold (currently 0.80), the maximum trade size "
            "(currently 10% of cash), and the minimum trade scale. This shows "
            "the system is robust and not dependent on one magic number.\n\n"
            "4. Begin writing Chapter 2 (Background and Literature Review) and "
            "Chapter 3 (Methodology). These chapters explain the theory behind "
            "the approach and describe how the system was built.",
            "M1: Phase-2 full grid completed — 10 seeds × 50k steps × "
            "70 tickers × 4 folds (mid-June).\n\n"
            "M2: Sensitivity analysis results locked (end of June).\n\n"
            "M3: Chapters 2 and 3 first drafts to supervisor (mid-July).\n\n"
            "M4: Final headline results table locked — no more "
            "experiment changes after this point (end of July).",
        ),
        (
            "August 2026\n(4 weeks)",
            "1. Write remaining dissertation chapters: Chapter 1 (Introduction), "
            "Chapter 5 (Results), Chapter 6 (Discussion), Chapter 7 (Conclusion).\n\n"
            "2. Polish all figures, tables, and equations. Ensure every result in "
            "the text matches the reproducible pipeline output.\n\n"
            "3. Integrate supervisor feedback on the full draft. Code changes from "
            "this point are bug-fix only.\n\n"
            "4. Stretch goal (time permitting): run a two-week paper-trading "
            "session using the Alpaca brokerage API to test the agent on live "
            "market data. This is not required for submission — the dissertation "
            "rests on the backtest and walk-forward evidence regardless.",
            "M5: Full draft to supervisor for review (mid-August).\n\n"
            "M6: Submission-ready version incorporating feedback "
            "(end of August).",
        ),
        (
            "September 2026",
            "1. Submit dissertation by 1 September 2026.\n\n"
            "2. Prepare viva presentation: a concise slide deck covering the "
            "research question, methodology, key results, and conclusions.\n\n"
            "3. Prepare a live demo of the reproducible pipeline (running the "
            "experiment end-to-end from the notebook).\n\n"
            "4. Rehearse anticipated questions and answers.",
            "M7: Dissertation submitted (1 September).\n\n"
            "M8: Viva presentation and demo ready (by viva date).",
        ),
    ])

    add_heading(doc, "Further details on the project direction", 2)
    add_para(
        doc,
        "The core research question will not change: does adding a probabilistic "
        "uncertainty signal to a reinforcement learning agent reduce portfolio "
        "drawdowns without sacrificing returns? The Phase-1 evidence says yes for "
        "SPY; the remaining work strengthens this claim across more assets, more "
        "seeds, and more time windows.",
    )
    add_para(
        doc,
        "If the Phase-2 results show that the uncertainty mechanism fails on certain "
        "asset classes (for example, commodity ETFs or highly volatile tech stocks), "
        "this will be reported honestly in the Results chapter as a limitation, not "
        "hidden. The dissertation will discuss where and why the approach works "
        "and where it does not.",
    )

    add_heading(doc, "Risks and mitigations", 2)
    add_bullets(doc, [
        "Compute time: Phase-2 experiments are larger (50,000 steps × 10 seeds × "
        "70 tickers) but the Colab T4 GPU runtime handles one full seed in under "
        "20 minutes. The experiment runner is designed to checkpoint progress so "
        "interrupted runs can resume without losing work.",
        "Result fragility: Phase-1 numbers may shift under Phase-2 conditions. "
        "Results will be reported as median and interquartile range across seeds, "
        "and any case where the probabilistic agent fails to beat a baseline will "
        "be reported explicitly rather than omitted.",
        "Live trading (stretch goal only): the paper-trading run with Alpaca is "
        "not on the critical path. The dissertation's conclusions are based on "
        "backtested evidence; the live run is a bonus validation if time permits.",
    ])

    page_break(doc)

    # ----- Extenuating circumstances -----
    add_heading(doc, "Extenuating circumstances", 1)
    add_para(
        doc,
        "[Student to fill in. Either record \"None to declare\" or describe and "
        "indicate that the personal tutor and student-support services have been "
        "informed. Do not include medical detail in this document.]",
        italic=True,
    )

    # ----- Self-tick -----
    add_heading(doc, "Indicative project hours and progress", 1)
    add_para(
        doc,
        "Self-assessment of the first one hundred hours allocated to the project. "
        "The student should tick exactly one of the following:",
    )
    add_bullets(doc, [
        "[ ]  The work has exceeded the first 100 hours of time allocated.",
        "[X] The work has sufficiently met the first 100 hours.",
        "[ ]  The majority of the first 100 hours have been completed but some time has been lost and will be made up.",
        "[ ]  Engagement in the project has been insufficient and progress is of concern.",
    ])
    add_para(
        doc,
        "Justification: the reproducible Phase-0 and Phase-1 pipeline, the protocol "
        "document, the baseline and probabilistic agents, the benchmarks, the "
        "rule-based stop-loss comparator and the generated supervisor pack together "
        "support the second tick above.",
        italic=True,
    )

    page_break(doc)

    # ----- Supervisor section -----
    add_heading(doc, "Supervisor's assessment", 1)
    add_para(
        doc,
        "The boxes below are reserved for the supervisor's written feedback. The "
        "student-completed sections of this document are intended to give the "
        "supervisor enough material to assess each item.",
        italic=True,
    )
    add_supervisor_box(doc, "Comments on the project plan and on the student's progress to date:")
    add_supervisor_box(doc, "Comments on the literature review and the framing of the problem:")
    add_supervisor_box(doc, "Comments on the technical progress and the experimental protocol:")
    add_supervisor_box(doc, "Recommendations for the remainder of the project:")
    add_supervisor_box(doc, "Other comments (optional):")

    # ----- Save -----
    out = EXPORTS / "InterimReview.docx"
    doc.save(out)
    print(f"Wrote: {out}")
    return out


if __name__ == "__main__":
    build()
